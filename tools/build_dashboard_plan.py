from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\Admin\Desktop\FITLINK\FitLink_Dashboard_Incremental_Plan.docx")
BLUE = "2E74B5"
DARK = "17324D"
PALE = "E8EEF5"
PALE_GREEN = "EAF5EE"
GREEN = "2E7D4F"
GRAY = "5F6B76"
LIGHT = "F5F7FA"


def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths, header_fill=PALE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, 10, True, DARK)
    mark_header_row(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_font(run, 9.5, False, "263746")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)
    return p


def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run("☐ ")
    set_font(r, 11, False, BLUE)
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=PALE_GREEN, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    set_font(r, 10.5, True, accent)
    r = p.add_run(text)
    set_font(r, 10.5, False, DARK)
    mark_header_row(table.rows[0])
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

# Running furniture
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("FITLINK  |  DASHBOARD DELIVERY GUIDE")
set_font(r, 8.5, True, GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("FitLink Kenya • Internal product planning • August 2026")
set_font(r, 8.5, False, GRAY)

# Editorial cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(112)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PRODUCT DELIVERY GUIDE")
set_font(r, 10, True, GREEN)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("FitLink Dashboard")
set_font(r, 30, True, DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
r = p.add_run("Incremental roadmap from a simple provider overview to a complete marketplace intelligence system")
set_font(r, 14, False, BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(74)
r = p.add_run("Trainer • Gym • Sports Academy • Wellness Centre • FitLink Admin")
set_font(r, 10.5, False, GRAY)
add_callout(doc, "Core principle", "Release the smallest dashboard that helps providers act today. Add deeper analytics only after the underlying events are consistently captured and trusted.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared 3 August 2026")
set_font(r, 10, True, GRAY)
doc.add_page_break()

doc.add_heading("1. Purpose and recommendation", level=1)
doc.add_paragraph("This guide translates FitLink's pricing promises into an executable dashboard plan. It defines the premium feature baseline, the data that must be captured, and an incremental delivery sequence that starts with a simple dashboard and grows without requiring a redesign of the entire product.")
add_callout(doc, "Recommended first release", "One provider dashboard with four KPI cards, an upcoming-bookings list, a new-leads list, subscription status, and five quick actions. Avoid charts, attribution, forecasting, recruitment, and complex admin reporting in the first release.")
doc.add_heading("Success criteria", level=2)
for x in [
    "A provider can understand today's business status in under 30 seconds.",
    "The dashboard exposes work that needs action: leads, bookings, payments, verification, and plan renewal.",
    "Every displayed metric has a documented source, definition, owner, and last-updated state.",
    "Premium analytics are added only after the corresponding tracking data passes quality checks.",
]: add_bullet(doc, x)

doc.add_heading("2. Premium feature baseline", level=1)
add_table(doc, ["Audience / plan", "Monthly price", "Premium promises"], [
    ("Trainer Professional", "KSh 3,000", "Featured listing; priority search; unlimited bookings; verified badge; analytics dashboard; marketing and priority support"),
    ("Trainer Premium", "KSh 5,000", "Homepage feature; unlimited leads; advanced analytics; personal website profile; account manager; custom campaigns"),
    ("Facility Premium", "KSh 10,000", "Featured placement; homepage promotion; unlimited inquiries; recruitment tools; event promotion; unlimited gallery; analytics; 24/7 support"),
], [1900, 1300, 6160])
doc.add_heading("Plan rule to resolve before development", level=2)
doc.add_paragraph("The pricing data lists each tier independently. Product and engineering should explicitly confirm whether Trainer Premium inherits all Trainer Professional capabilities. This guide assumes inheritance because it is the clearest customer expectation and the simplest access-control model.")
doc.add_heading("Shared provider capabilities", level=2)
for x in [
    "Trainer: profile, M-Pesa payment support, booking management, messaging, certificate uploads, and mobile presentation.",
    "Facility: profile, membership management, trainer directory, events, photo gallery, and bookings.",
    "Commercial model: 1–2% commission on the first successful client-provider transaction; no repeat commission for that pair.",
]: add_bullet(doc, x)

doc.add_page_break()
doc.add_heading("3. Data foundation", level=1)
doc.add_paragraph("The dashboard should distinguish operational records from behavioral events and derived metrics. Operational records support daily work. Behavioral events explain discovery and conversion. Derived metrics summarize those records using documented formulas.")
add_table(doc, ["Data domain", "Core records", "Dashboard use"], [
    ("Identity & profile", "Users, providers, roles, approval, plan, services, documents, photos", "Access control, profile completion, verification, subscription"),
    ("Leads & inquiries", "Source, service, message, owner, status, response time, outcome", "Lead inbox, response SLA, conversion pipeline"),
    ("Bookings", "Client, provider, service, schedule, amount, status, attendance, cancellation", "Calendar, operations, revenue, retention"),
    ("Payments", "Method, reference, amount, verification, refund, subscription period", "Reconciliation, billing, outstanding actions"),
    ("Discovery events", "Search impression, result rank, profile view, CTA click, campaign tags", "Visibility, funnels, premium placement value"),
    ("Referrals & commissions", "Permanent pair lock, first booking, rate, amount, settlement", "Commission ledger and repeat-client protection"),
    ("Reputation", "Review, rating, verified booking, response, moderation", "Rating trend and response work"),
    ("Facility extensions", "Memberships, vacancies, applicants, events, attendance", "Facility operations and premium tools"),
], [1700, 4560, 3100])

doc.add_heading("Minimum event dictionary", level=2)
for x in [
    "search_performed — query, filters, location, session and timestamp",
    "provider_impression — provider, placement, rank, page and campaign",
    "provider_profile_viewed — provider, source, session and user when known",
    "provider_cta_clicked — action type: inquiry, WhatsApp, book, call or event",
    "lead_created — provider, client, source, requested service and campaign",
    "booking_created / confirmed / completed / cancelled / no_show",
    "payment_submitted / verified / failed / refunded",
    "review_submitted / provider_responded",
]: add_bullet(doc, x)
add_callout(doc, "Privacy rule", "Do not place full phone numbers, transaction references, private documents, or client health details in summary widgets. Use role-based access and reveal sensitive fields only where the workflow requires them.", fill="FFF4E5", accent="9A6700")

doc.add_page_break()
doc.add_heading("4. Incremental delivery roadmap", level=1)
add_table(doc, ["Phase", "Outcome", "Complexity", "Primary users"], [
    ("1. Simple overview", "See and act on today's work", "Low", "All providers"),
    ("2. Operations", "Manage leads, bookings, clients and payments", "Medium", "All providers"),
    ("3. Basic analytics", "Understand visibility, conversion and revenue", "Medium", "Professional / Premium"),
    ("4. Premium growth", "Measure campaigns, cohorts and promotion value", "High", "Premium providers"),
    ("5. Facility tools", "Run memberships, recruitment and events", "High", "Premium facilities"),
    ("6. Admin command centre", "Operate marketplace, billing and trust", "High", "FitLink team"),
], [1700, 4200, 1300, 2160])

doc.add_heading("Phase 1 — Simple provider overview", level=1)
doc.add_paragraph("Goal: give every provider a useful home screen without requiring mature analytics infrastructure. This phase should work with the provider, booking, lead, payment, review and subscription records already planned for the platform.")
doc.add_heading("Screen contents", level=2)
for x in [
    "Four KPI cards: upcoming bookings, new leads, pending payment value, and current rating.",
    "Upcoming bookings: next five records with time, client, service, status and View action.",
    "New leads: newest five inquiries with age, requested service, source and Respond action.",
    "Subscription panel: plan, payment status, renewal/expiry date and upgrade or resolve-payment action.",
    "Profile status: approval, verification and a short profile-completion checklist.",
    "Quick actions: manage availability, add booking, edit profile, add event, contact support.",
]: add_bullet(doc, x)
doc.add_heading("Deliberate exclusions", level=2)
for x in [
    "No line, bar, pie or funnel charts.",
    "No marketing attribution, cohorts, forecasting or benchmarks.",
    "No configurable widgets or saved dashboard layouts.",
    "No recruitment, complex memberships or cross-provider rollups.",
]: add_bullet(doc, x)
doc.add_heading("Phase 1 acceptance checklist", level=2)
for x in [
    "Role and provider ownership are enforced on every query.",
    "All cards show an empty, loading, success and error state.",
    "Every count links to the corresponding operational list.",
    "The view is usable at 360 px width without horizontal scrolling.",
    "Amounts use KSh consistently; times use Africa/Nairobi by default.",
    "A new provider sees onboarding guidance instead of zero-filled noise.",
    "No metric is calculated differently between dashboard and detail screens.",
]: add_check(doc, x)

doc.add_page_break()
doc.add_heading("Phase 2 — Operational workspace", level=1)
doc.add_paragraph("Goal: turn the overview into a complete daily-work system while keeping the overview itself simple.")
for title, items in [
    ("Leads", ["Status pipeline: new, contacted, qualified, converted, lost", "Owner, follow-up date, notes, response-time indicator and loss reason", "Search, filters, sorting, bulk assignment for facility teams"]),
    ("Bookings", ["Calendar and list views", "Confirm, reschedule, cancel, mark attendance and complete", "Conflict warnings, client history, payment and refund state"]),
    ("Clients", ["Client summary, first/repeat flag and booking history", "Communication preferences and consent", "Provider-scoped notes with restricted access"]),
    ("Payments", ["Pending-verification queue and transaction references", "Receipts, refunds, discrepancies and reconciliation state", "Subscription invoice and renewal history"]),
]:
    doc.add_heading(title, level=2)
    for x in items: add_bullet(doc, x)

doc.add_heading("Phase 3 — Basic analytics", level=1)
doc.add_paragraph("Goal: fulfill the analytics-dashboard promise with understandable, trustworthy metrics before introducing advanced attribution.")
for x in [
    "Date selector: 7, 30 and 90 days plus custom range.",
    "KPI trend: profile views, leads, bookings, booking value and conversion rate.",
    "Simple funnel: impressions → profile views → leads → bookings.",
    "Bookings and revenue trend by day/week.",
    "Top services and lead sources.",
    "New versus returning clients, cancellation rate and no-show rate.",
    "Metric tooltip, previous-period comparison and last-updated timestamp.",
]: add_bullet(doc, x)
add_callout(doc, "Analytics gate", "Do not release a chart until its source events have at least one complete reporting period, duplicate-event controls, known timezone handling, and reconciliation against transactional totals.")

doc.add_page_break()
doc.add_heading("Phase 4 — Premium growth analytics", level=1)
doc.add_paragraph("Goal: prove the business value of homepage promotion, featured search placement and custom campaigns.")
for x in [
    "Organic versus promoted impressions, profile visits, leads and bookings.",
    "Average search position and visibility by category, county and device.",
    "Campaign performance: impressions, clicks, leads, bookings and attributed value.",
    "Lead and client cohorts by acquisition month, service and source.",
    "Repeat-booking rate, estimated lifetime value and retention windows.",
    "Downloadable CSV/PDF reports and scheduled summaries.",
    "Account-manager notes, campaign approvals and recommendations.",
]: add_bullet(doc, x)

doc.add_heading("Phase 5 — Facility premium modules", level=1)
for title, items in [
    ("Memberships", ["Inquiry-to-member pipeline", "Plan, start, renewal, pause and expiry", "Branch, service and membership revenue views"]),
    ("Recruitment", ["Vacancies, applicants and qualifications", "Pipeline stages, interview dates and hiring outcome", "Restricted recruiter permissions and document access"]),
    ("Events", ["Create, publish and promote events", "Capacity, registrations, attendance and waitlist", "Event revenue and promotion conversion"]),
    ("Multi-location", ["Branch filter and branch-level permissions", "Location comparison and roll-up totals", "Shared versus location-specific trainers and services"]),
]:
    doc.add_heading(title, level=2)
    for x in items: add_bullet(doc, x)

doc.add_heading("Phase 6 — FitLink admin command centre", level=1)
for x in [
    "Marketplace overview: active providers, clients, bookings, value and conversion.",
    "Provider approval and document-verification queue.",
    "Payment verification, subscriptions, delinquency and refunds.",
    "Referral locks, commission ledger, settlement and discrepancy workflow.",
    "Promotion inventory, featured-placement scheduling and campaign approvals.",
    "Review moderation, support SLA, fraud alerts and immutable audit log.",
]: add_bullet(doc, x)

doc.add_page_break()
doc.add_heading("5. Dashboard design checklist", level=1)
sections = {
    "Product and metric rules": [
        "Confirm tier inheritance and access-control rules.", "Define lead, inquiry, booking, successful booking and active membership.", "Document formulas, refresh rate, owner and source for every KPI.", "Define plan-limit reset and over-limit behavior.", "Define featured-placement ranking and fair-rotation rules."],
    "Information architecture": [
        "Provide role-specific navigation for trainer, facility and admin users.", "Keep the overview focused on status and next actions.", "Use global provider/location/date filters only where relevant.", "Show current plan and urgent account state persistently.", "Design clear empty, loading, error and permission-denied states."],
    "Interaction and accessibility": [
        "Ensure keyboard navigation and visible focus states.", "Use text labels alongside color and icons.", "Meet WCAG AA contrast for text and interactive controls.", "Make tables responsive with card/list alternatives on small screens.", "Confirm actions with clear success/error feedback and undo where practical."],
    "Security and privacy": [
        "Enforce provider ownership and role permissions server-side.", "Mask phone and transaction details in summary views.", "Keep certificates and verification documents private.", "Record audit events for approval, payment and commission changes.", "Support data retention, export and deletion policies."],
    "Data quality": [
        "Use stable IDs and server timestamps.", "Make event ingestion idempotent and deduplicate retries.", "Attach source/campaign fields to lead and booking creation.", "Exclude bots, test accounts and internal traffic from analytics.", "Reconcile analytical totals with booking and payment records."],
}
for title, items in sections.items():
    doc.add_heading(title, level=2)
    for x in items: add_check(doc, x)

doc.add_heading("6. Suggested implementation order", level=1)
steps = [
    "Write the metric dictionary and plan-entitlement matrix.",
    "Finalize Firestore records and indexes for providers, leads, bookings, payments, subscriptions and reviews.",
    "Build shared dashboard shell, navigation, permissions and responsive layout.",
    "Deliver Phase 1 overview using transactional data only.",
    "Add Phase 2 detail screens and operational actions.",
    "Instrument discovery and conversion events; validate them for one full reporting period.",
    "Release basic analytics, then premium attribution and facility modules.",
    "Build the admin command centre once provider workflows and data definitions have stabilized.",
]
for x in steps: add_number(doc, x)

doc.add_heading("7. Phase 1 screen blueprint", level=1)
add_table(doc, ["Area", "Component", "Required fields / behavior"], [
    ("Header", "Greeting + account state", "Provider name, plan badge, verification state, notification entry"),
    ("KPI row", "4 cards", "Upcoming bookings; new leads; pending payment value; rating"),
    ("Primary", "Upcoming bookings", "Next five; date/time; client; service; status; view action"),
    ("Primary", "New leads", "Newest five; age; service; source; respond action"),
    ("Secondary", "Subscription", "Plan; status; renewal/expiry; billing action"),
    ("Secondary", "Profile progress", "Approval; verification; missing profile items"),
    ("Actions", "Quick-action group", "Availability; add booking; edit profile; add event; support"),
], [1350, 2200, 5810])

doc.add_heading("Definition of done for the simple dashboard", level=2)
for x in [
    "Test data demonstrates a new provider, active provider, premium provider and facility.",
    "KPI values match source queries in automated tests.",
    "All visible actions navigate or execute successfully.",
    "Mobile, tablet and desktop layouts have been visually reviewed.",
    "Access-control tests prevent cross-provider data exposure.",
    "Performance is acceptable on a typical mobile connection.",
    "Product owner approves metric labels, empty states and upgrade messaging.",
]: add_check(doc, x)

doc.add_heading("8. Source references", level=1)
doc.add_paragraph("This plan is based on the current FitLink pricing and implementation files:")
for x in [
    r"src/data/pricing.js — plan names, prices, feature promises and included capabilities",
    r"src/pages/Pricing.jsx — customer-facing pricing presentation and commission copy",
    r"src/lib/registrations.js — provider, plan, payment, photo and document records",
    r"src/components/PaymentModal.jsx — current manual Pochi/M-Pesa confirmation data",
    r"scripts/seed.js — production-shaped users, providers, bookings and plan data",
    r"scripts/commissions.js — permanent first-client referral lock and commission ledger",
]: add_bullet(doc, x)

doc.core_properties.title = "FitLink Dashboard Incremental Plan"
doc.core_properties.subject = "Incremental dashboard design and delivery checklist"
doc.core_properties.author = "FitLink Product Planning"
doc.core_properties.keywords = "FitLink, dashboard, analytics, premium, roadmap"
doc.save(OUT)
print(OUT)
